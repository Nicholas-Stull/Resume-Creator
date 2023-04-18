from docx import Document
from docx.shared import Inches,Pt
import re,qrcode
from docx.enum.text import WD_ALIGN_PARAGRAPH

# creating document object
document = Document()

#name
name = input("Enter your name: ")
name = ' '.join(elem.capitalize() for elem in name.split())
name_heading = document.add_heading(name,0)
name_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

#job subheading
job = input("Enter your title: ")
job = ' '.join(elem.capitalize() for elem in job.split())
job_heading = document.add_heading(job,1)
job_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

#city and state
city = input("Enter your City, State: ")
city = ' '.join(elem.capitalize() for elem in city.split())

#phone
phone = input("Enter your phone number: ")
phone = re.sub('[^0-9]', '', phone)
formating_phone = f"({phone[:3]}) {phone[3:6]}-{phone[6:]}"

#email
email = input("Enter your email: ")
email = ' '.join(elem.capitalize() for elem in email.split())

#qr code
data = input("Enter the URL you want to encode to the qr code: ")
img = qrcode.make(data)
img.save('qr.png')

#contact details
contact_details = document.add_paragraph()
contact_details.add_run(city + '\n').bold = True
contact_details.add_run(formating_phone + '\n').bold = True
contact_details.add_run(email + '\n').bold = True
contact_details.alignment = WD_ALIGN_PARAGRAPH.CENTER

#summary
summary_heading = document.add_heading('Summary', level=3)
document.add_paragraph(input('Tell us a little about yourself: '))

# work experience
experience_heading = document.add_heading('Work Experience', level=3)
p = document.add_paragraph()
print("Let's gather your work experience:")
company = input('Please enter company name: ')
company_up = ' '.join(elem.capitalize() for elem in company.split())
title = input('Please enter the title of your position at this job: ')
from_date = input('Date Started mm/dd/yyyy: ')
to_date = input('End Date mm/dd/yyyy: ')

company_name = p.add_run(company_up)
company_name.bold = True
company_name.underline = True
p.add_run(f' \n')

p.add_run(f'{title} \n').bold = True
p.add_run(f'{from_date} - {to_date} \n').italic = True

experience_details = input(f'Describe your job experience at {company_up}: ')
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input('Do you have any other experience? Yes or No: ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Please enter company name: ')
        company_up = ' '.join(elem.capitalize() for elem in company.split())
        title = input('Please enter the title of your position at this job: ')
        from_date = input('Date Started mm/dd/yyyy: ')
        to_date = input('End Date mm/dd/yyyy: ')

        company_name = p.add_run(company_up)
        company_name.bold = True
        company_name.underline = True
        p.add_run(f' \n')

        p.add_run(f'{title} \n').bold = True
        p.add_run(f'{from_date} - {to_date} \n').italic = True

        experience_details = input(f'Describe your job experience at {company_up}: ')
        p.add_run(experience_details)
    else:
        break

# Credentials: Degrees & IT Certifications
# Add degrees
document.add_heading('Degrees',level=3)
p = document.add_paragraph()
degree = input('Please enter your highest degree: ')
place = input('Please enter the place of your degree: ')
from_date = input('Date recieved mm/dd/yyyy: ')
p.add_run(f'{degree} \n').bold = True
p.add_run(f'{place} \n').bold = True
p.add_run(f'{from_date} \n').italic = True
while True:
    has_more_degree = input('Do you have any other Degree? Yes or No: ')
    if has_more_degree.lower() == 'yes':
        p = document.add_paragraph()
        degree2 = input('Please enter degree: ')
        place2 = input('Please enter the place of your degree: ')
        from_date2 = input('Date recieved mm/dd/yyyy: ')
        p.add_run(f'{degree2} \n').bold = True
        p.add_run(f'{place2} \n').bold = True
        p.add_run(f'{from_date2} \n').italic = True
    else:
        break


# Add certifications
document.add_heading('Certifications',level=3)
p = document.add_paragraph()
certification = input('Please enter your IT certification: ')
p.add_run(f'{certification} \n').bold = True


# Skills
skills_heading = document.add_heading('Skills', level=3)
#skills = document.add_paragraph()
skills = []
while len(skills) < 12:
    skill = input("Enter a skill (press enter to add another): ").capitalize()
    if skill:
        skills.append(skill)
    else:
        break
#skills_list = document.add_paragraph(style='List Bullet')
for skill in skills:
    skills_list = document.add_paragraph(style='List Bullet')
    skills_list.add_run(skill)
docu_name = ('_'.join(elem.capitalize() for elem in name.split()) + "_Resume.docx")
document.save(docu_name)
